# 增强版文档解析服务
import os
import PyPDF2
from docx import Document as DocxDocument
from typing import Optional, List, Dict, Any
import json

class DocumentParser:
    """增强版文档解析器 - 支持更多文件类型"""
    
    SUPPORTED_TYPES = {
        'pdf': 'PDF文档',
        'docx': 'Word文档',
        'doc': 'Word文档',
        'txt': '文本文件',
        'md': 'Markdown文档',
        'json': 'JSON文件',
        'csv': 'CSV文件',
        'xlsx': 'Excel表格',
        'xls': 'Excel表格'
    }
    
    def parse(self, file_path: str, filename: str) -> Dict[str, Any]:
        """根据文件类型解析文档，返回结构化数据"""
        ext = filename.split('.')[-1].lower()
        
        result = {
            'filename': filename,
            'file_type': ext,
            'content': '',
            'metadata': {},
            'success': True,
            'error': None
        }
        
        try:
            if ext == 'pdf':
                result['content'] = self._parse_pdf(file_path)
            elif ext in ['docx', 'doc']:
                result['content'] = self._parse_docx(file_path)
            elif ext == 'txt':
                result['content'] = self._parse_txt(file_path)
            elif ext == 'md':
                result['content'] = self._parse_markdown(file_path)
            elif ext == 'json':
                result['content'] = self._parse_json(file_path)
            elif ext == 'csv':
                result['content'] = self._parse_csv(file_path)
            elif ext in ['xlsx', 'xls']:
                result['content'] = self._parse_excel(file_path)
            else:
                # 其他类型尝试作为文本读取
                result['content'] = self._parse_txt(file_path)
                result['metadata']['warning'] = f'未知文件类型: {ext}，尝试作为文本读取'
            
            # 提取通用元数据
            result['metadata'].update({
                'char_count': len(result['content']),
                'line_count': len(result['content'].split('\n')),
                'word_count': len(result['content'].split())
            })
            
        except Exception as e:
            result['success'] = False
            result['error'] = str(e)
            result['content'] = f"[解析错误: {str(e)}]"
        
        return result
    
    def _parse_pdf(self, file_path: str) -> str:
        """解析PDF文件"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata = pdf_reader.metadata
                num_pages = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- 第{page_num + 1}页 ---\n{page_text}\n"
                
                # 保存PDF元数据
                self._last_metadata = {
                    'total_pages': num_pages,
                    'pdf_metadata': {k: str(v) for k, v in metadata.items()} if metadata else {}
                }
        except Exception as e:
            text = f"[PDF解析错误: {str(e)}]"
        return text.strip()
    
    def _parse_docx(self, file_path: str) -> str:
        """解析Word文件"""
        text = ""
        try:
            doc = DocxDocument(file_path)
            
            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # 提取表格内容
            for table_idx, table in enumerate(doc.tables):
                text += f"\n--- 表格{table_idx + 1} ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += row_text + "\n"
            
            # 保存元数据
            self._last_metadata = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }
        except Exception as e:
            text = f"[Word解析错误: {str(e)}]"
        return text.strip()
    
    def _parse_txt(self, file_path: str) -> str:
        """解析文本文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        return "[文本解析错误: 无法解码文件]"
    
    def _parse_markdown(self, file_path: str) -> str:
        """解析Markdown文件"""
        content = self._parse_txt(file_path)
        # 可以在这里添加Markdown特定的处理
        return content
    
    def _parse_json(self, file_path: str) -> str:
        """解析JSON文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
                return json.dumps(data, ensure_ascii=False, indent=2)
        except Exception as e:
            return f"[JSON解析错误: {str(e)}]"
    
    def _parse_csv(self, file_path: str) -> str:
        """解析CSV文件"""
        try:
            import csv
            text = ""
            with open(file_path, 'r', encoding='utf-8') as file:
                reader = csv.reader(file)
                for row in reader:
                    text += " | ".join(row) + "\n"
            return text
        except Exception as e:
            return f"[CSV解析错误: {str(e)}]"
    
    def _parse_excel(self, file_path: str) -> str:
        """解析Excel文件"""
        try:
            import pandas as pd
            text = ""
            # 读取所有sheet
            xls = pd.ExcelFile(file_path)
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                text += f"\n--- Sheet: {sheet_name} ---\n"
                text += df.to_string(index=False)
                text += "\n"
            return text
        except ImportError:
            return "[Excel解析错误: 请安装pandas和openpyxl]"
        except Exception as e:
            return f"[Excel解析错误: {str(e)}]"
    
    def get_supported_types(self) -> List[str]:
        """获取支持的文件类型列表"""
        return list(self.SUPPORTED_TYPES.keys())
    
    def is_supported(self, filename: str) -> bool:
        """检查文件类型是否支持"""
        ext = filename.split('.')[-1].lower()
        return ext in self.SUPPORTED_TYPES


# 文档分块处理器
class DocumentChunker:
    """文档分块处理器 - 用于RAG检索优化"""
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def chunk_text(self, text: str) -> List[Dict[str, Any]]:
        """将文本分割成块"""
        chunks = []
        
        # 按段落分割
        paragraphs = text.split('\n\n')
        current_chunk = ""
        chunk_index = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # 如果当前块加上新段落超过限制
            if len(current_chunk) + len(para) > self.chunk_size:
                if current_chunk:
                    chunks.append({
                        'index': chunk_index,
                        'content': current_chunk.strip(),
                        'char_count': len(current_chunk)
                    })
                    chunk_index += 1
                    # 保留重叠部分
                    current_chunk = current_chunk[-self.chunk_overlap:] if len(current_chunk) > self.chunk_overlap else ""
                current_chunk += para + "\n\n"
            else:
                current_chunk += para + "\n\n"
        
        # 添加最后一个块
        if current_chunk.strip():
            chunks.append({
                'index': chunk_index,
                'content': current_chunk.strip(),
                'char_count': len(current_chunk)
            })
        
        return chunks
    
    def chunk_by_sentences(self, text: str) -> List[Dict[str, Any]]:
        """按句子分割文本"""
        import re
        sentences = re.split(r'[。！？.!?]\s*', text)
        
        chunks = []
        current_chunk = ""
        chunk_index = 0
        
        for sent in sentences:
            sent = sent.strip()
            if not sent:
                continue
            
            if len(current_chunk) + len(sent) > self.chunk_size:
                if current_chunk:
                    chunks.append({
                        'index': chunk_index,
                        'content': current_chunk.strip(),
                        'char_count': len(current_chunk)
                    })
                    chunk_index += 1
                current_chunk = sent + "。"
            else:
                current_chunk += sent + "。"
        
        if current_chunk.strip():
            chunks.append({
                'index': chunk_index,
                'content': current_chunk.strip(),
                'char_count': len(current_chunk)
            })
        
        return chunks
